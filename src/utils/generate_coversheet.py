import sys
from docx import Document
import os

def generate_coversheet(stock_number, manufacturer, model, year, web_width, colors, die_stations, description, dropbox_url, owner_company, owner_name, owner_phone, owner_email, price, buy_price, notes):
    doc = Document()
    doc.add_heading('Coversheet', 0)

    table = doc.add_table(rows=0, cols=2)
    data = [
        ('Stock Number', stock_number),
        ('Category', 'Press'),
        ('Manufacturer', manufacturer),
        ('Model', model),
        ('Year', year),
        ('Web Width', web_width),
        ('Colors', colors),
        ('Die Stations', die_stations),
        ('Description', description),
        ('Dropbox URL', dropbox_url),
        ('Owner Company', owner_company),
        ('Owner Name', owner_name),
        ('Owner Phone', owner_phone),
        ('Owner Email', owner_email),
        ('Price', price),
        ('Buy Price', buy_price),
        ('Additional Notes', notes)
    ]

    for label, value in data:
        row_cells = table.add_row().cells
        row_cells[0].text = label
        row_cells[1].text = value if value else ''

    filename = f'{stock_number} Coversheet.docx'
    save_path = os.path.join(os.getcwd(), filename)
    doc.save(save_path)

if __name__ == "__main__":
    generate_coversheet(*sys.argv[1:])

